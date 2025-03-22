from RPA.Browser.Selenium import Selenium
from openpyxl import load_workbook
import os
from RPA.Email.ImapSmtp import ImapSmtp
import urllib.request
from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from dotenv import load_dotenv

class SportsShopSite:
    def __init__(self, url, download_directory):
        self.url = url
        self.download_directory= download_directory
        self.browser = Selenium()
        self.browser.set_download_directory(self.download_directory)
        self.browser.open_chrome_browser(url)
    
    def extract_new_products_data(self):
        self.browser.wait_until_element_is_visible("//html/body/center/table[2]/tbody")
        sports_equipment = self.browser.find_elements("//html/body/center/table[2]/tbody/tr/td")
        new_equipments = []
        for equip in sports_equipment:
            print(equip.text)
        for i, equip in enumerate(sports_equipment):
            element = self.browser.find_elements(f"//*[@id='Sport{i+1}']/span")
            if element:
                new_equipments.append(equip)
                image_xpath = f"//html/body/center/table[2]/tbody/tr[{(i+2)//2 }]/td[{(i)%2 +1}]/div/div/div/table/tbody/tr[1]/td[1]/img"
                print(image_xpath)
                href_image = self.browser.get_element_attribute(image_xpath, "src")
                print(href_image)
                image_name = href_image.split("/")[-1]
                urllib.request.urlretrieve(href_image, f"{self.download_directory}/{image_name}")
        new_equipments_data = [i.text for i in new_equipments]
        return new_equipments_data
        
    def extract_school_data(self):
        self.browser.go_to("https://botsdna.com/sportshop/schools.html")
        self.browser.wait_until_element_is_visible("//*[@id='courts']/tbody")
        rows = self.browser.find_elements("//*[@id='courts']/tbody/tr")[1:]
        school_data = []
        for i, elements in enumerate(rows):
            school_dict = {}
            school_dict["school_code"] = self.browser.get_text(f"//*[@id='courts']/tbody/tr[{i+2}]/td[1]")
            school_dict["school_name"] = self.browser.get_text(f"//*[@id='courts']/tbody/tr[{i+2}]/td[2]")
            school_dict["student_strength"] = int(self.browser.get_text(f"//*[@id='courts']/tbody/tr[{i+2}]/td[3]").strip())
            school_data.append(school_dict)
        return school_data

    def download_excel_file(self):
        self.browser.click_element("//a[contains(@href, 'EmailsDatabase.xlsx')]")
        self.waiting(f"{self.download_directory}/EmailsDatabase.xlsx")
    
    def download_email_attachment_template(self):
        self.browser.click_element("//a[contains(@href, 'SportsTemplet.docx')]")
        self.waiting(f"{self.download_directory}/SportsTemplet.docx")
    
    def waiting(self, file_path):
        while not os.path.exists(file_path):
            continue
        return

class ExcelHandling:
    def __init__(self, excel_file_path):
        self.file_path = excel_file_path
    
    def get_email_ids(self):
        wb = load_workbook(filename= self.file_path)
        ws = wb.worksheets[0]
        email_id_data = {}
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column, values_only=True):
            school_code = row[0]
            school_emails = []
            for email in row[1:]:
                if email:
                    school_emails.append(email)
            email_id_data[school_code] = school_emails
        return email_id_data
    
    def get_price_offer(self):
        wb = load_workbook(filename= self.file_path)
        ws = wb.worksheets[1]
        offers = []
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column, values_only=True):
            start, end, offer_percent = row
            offers.append([start, end, offer_percent])
        return offers

class Mail:
    def __init__(self,absolute_pdf_path):
        self.absolute_pdf_path = absolute_pdf_path
        self.mail = ImapSmtp()

        self.gmail_account =  os.getenv("GMAIL_ACCOUNT")
        self.gmail_password = os.getenv("GMAIL_PASSWORD")

        self.mail.authorize(
            account=self.gmail_account,
            password=self.gmail_password,
            smtp_server="smtp.gmail.com",
            smtp_port=587,
        )
    
    def send_mail(self,recipients,school_name):
        data = """Dear Customer,\nIt is my great pleasure to offer you the best price in the latest stock,\nPlease take a moment to review the attached sport products & get back with order\nShould you have any further questions, please do not hesitate to contact me.\nSincerely,\nSudheer Nimmagadda\nbotsDNA Sports Shop\n+91 9705435277"""
        file_to_send = f"{self.absolute_pdf_path}/{school_name}-botsDNA sports Quote.pdf"

        self.mail.send_message(
            recipients=recipients,
            sender= self.gmail_account,
            subject= f"{school_name} - New Sports products arrived with Great offers",
            body=data,
            attachments=file_to_send
        )
        print("mail_sent")
        return

class CreatingPdfs:
    def __init__(self, pdf_directory, template_file_path):
        self.pdf_directory = pdf_directory
        self.template_file_path = template_file_path
        if not os.path.exists(self.pdf_directory):
            os.makedirs(self.pdf_directory)
    
    def add_data_to_docx(self, school_name, new_equipments_data, offer_percent):
        document = Document(self.template_file_path)
        intro_para = document.add_paragraph(f"Your school {school_name} had {offer_percent} % off on all products")
        intro_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        for i, product in enumerate(new_equipments_data):
            lines = product.split("\n")[1:-1]
            product_name = lines[0].split(":")[-1].strip()
            product_code = lines[1].split(":")[-1].strip()
            price = int(lines[2].split(":")[-1].strip()[:-2])
            effective_price = int(price - (price * (offer_percent/100)))
            product_para = document.add_paragraph(f"\n#{i + 1} \nPRODUCT NAME: {product_name} \nPRODUCT CODE: {product_code} \nUNIT PRICE: R{price} & OFFER PRICE: R{effective_price}"
            )
            product_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_path = f"./downloads/{product_code}.jpg"
            my_image = document.add_picture(image_path)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        file_name = f"{self.pdf_directory}/{school_name}-botsDNA sports Quote"
        document.save(f"{file_name}.docx")
        convert(f"{file_name}.docx", f"{file_name}.pdf")
        
class Process:
    def __init__(self, url, download_directory, pdf_directory,absolute_pdf_path):
        self.url = url
        self.download_directory = download_directory
        self.pdf_directory = pdf_directory
        self.absolute_pdf_path = absolute_pdf_path
    
    def send_sports_products_data_to_schools(self):
        website = SportsShopSite(self.url, self.download_directory)
        new_equipments_data= website.extract_new_products_data()
        website.download_excel_file()
        website.download_email_attachment_template()
        school_data = website.extract_school_data()

        excel = ExcelHandling(f"{self.download_directory}/EmailsDatabase.xlsx")
        school_email_id_data = excel.get_email_ids()
        offers = excel.get_price_offer()

        mail = Mail(self.absolute_pdf_path)
        pdf = CreatingPdfs(self.pdf_directory, f"{self.download_directory}/SportsTemplet.docx")

        for schools in school_data:
            school_code = schools["school_code"]
            school_name = schools["school_name"]
            student_strength = schools["student_strength"]
            
            offer = 0
            for start, end, offer_percent in offers:
                if start <= student_strength <= end:
                    offer = offer_percent
                    break
            if offer == 0 and student_strength> offers[-1][1]:
                offer = 25
            
            pdf.add_data_to_docx(school_name, new_equipments_data, offer)

            recipents = []
            for schl_code in school_email_id_data.keys():
                if schl_code == school_code:
                    recipents = school_email_id_data[schl_code]
                    break
            if recipents:
                mail.send_mail(recipents, school_name)
            

def main():
    load_dotenv()
    URL = "https://botsdna.com/sportshop/index.html"
    DOWNLOAD_DIRECTORY = "./downloads"
    PDF_DIRECTORY = "./pdfs"
    ABSOLUTE_PDF_PATH = os.getenv("ABSOLUTE_PDF_PATH")
    process = Process(URL, DOWNLOAD_DIRECTORY, PDF_DIRECTORY, ABSOLUTE_PDF_PATH)
    process.send_sports_products_data_to_schools()

main()